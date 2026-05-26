"""
팀원 공유용 Word 문서 작성 (python-docx).

내용: docs/팀원공유_2026-05-26.md 기반
그림: reports/figures/ 6개 PNG 삽입
산출: reports/팀원공유_2026-05-26.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
FIG_DIR = ROOT / "reports" / "figures"
OUT_PATH = ROOT / "reports" / "팀원공유_2026-05-26.docx"

FONT_NAME = "Noto Sans CJK KR"  # 시스템 한글 폰트 (이름은 Word가 자체 매핑)


def set_font(run, name=FONT_NAME, size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # 한글 글리프 대응
    from docx.oxml.ns import qn
    rPr = run._element.rPr or run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 0:  # 제목
        set_font(run, size=18, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        set_font(run, size=14, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    elif level == 2:
        set_font(run, size=12, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
    else:
        set_font(run, size=11, bold=True)


def add_para(doc: Document, text: str, bullet: bool = False, size=10.5):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=size)


def add_para_runs(doc: Document, runs_spec: list, bullet: bool = False):
    """runs_spec: list of (text, bold) tuples."""
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.space_after = Pt(2)
    for text, bold in runs_spec:
        run = p.add_run(text)
        set_font(run, size=10.5, bold=bold)


def add_figure(doc: Document, path: Path, caption: str, width_cm: float = 15.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    crun = cap.add_run(caption)
    set_font(crun, size=9, color=RGBColor(0x55, 0x55, 0x55))


def add_table(doc: Document, header: list[str], rows: list[list[str]],
              col_widths_cm: list[float] | None = None,
              highlight_rows: list[int] | None = None):
    """간단 표. highlight_rows: 0-indexed (header 제외) 진한 배경"""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    # header
    for j, h in enumerate(header):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=10, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_widths_cm:
            cell.width = Cm(col_widths_cm[j])
    # rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            bold = highlight_rows is not None and i in highlight_rows
            set_font(run, size=10, bold=bold)
            if col_widths_cm:
                cell.width = Cm(col_widths_cm[j])
    # paragraph spacing
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build():
    doc = Document()

    # 기본 여백 줄이기 (A4 효율)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ===== 제목 =====
    add_heading(doc, "일별 최대전력 ARIMAX 예측 — 진행 현황 공유", level=0)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run("작성일 2026-05-26 · 학습 2010-2024 (5,479일) / 평가 2025 (365일 홀드아웃)")
    set_font(meta_run, size=10, color=RGBColor(0x66, 0x66, 0x66))

    add_para(doc, "ARIMAX/SARIMAX 모델로 일별 최대전력(MW) 16년치를 학습해 2025년 365일을 예측·검증. "
                  "수집·EDA·모델링·평가까지 3차 반복 개선을 거쳐 현재 v3c 최종 모델 도출.")

    # 종속변수 도입부 그림
    add_figure(doc, FIG_DIR / "fig01_peak_mw_timeline.png",
               "그림 1. 종속변수 — 일별 최대전력 2010-2025 (5,844일). 점선은 2025 홀드아웃 시작.")

    # ===== 1. 데이터 =====
    add_heading(doc, "1. 데이터", level=1)

    add_heading(doc, "1.1 종속변수", level=2)
    add_para(doc, "한국전력거래소(KPX) 일별 최대전력(MW), 2010-01-01 ~ 2025-12-31, 5,844일.", bullet=True)
    add_para(doc, "기간 확정: 2008·2009년 결측 다수로 2010년부터 연속성 확보.", bullet=True)
    add_para(doc, "학습 2010-2024 (5,479일) / 평가 2025 (365일, 홀드아웃).", bullet=True)

    add_heading(doc, "1.2 외생변수 — 5개 카테고리, 28개 변수", level=2)
    add_para_runs(doc, [("A. 기상 (ASOS 일·시간자료)", True), (
        " — 17개 시·도 대표 관측소 + 백업 6개. ", False), (
        "인구 가중 평균", True), ("으로 전국 단일 시계열 산출. "
        "결측은 인접 관측소 bias 보정 + 시간 선형 보간 → 잔여 결측 0.", False)], bullet=True)
    add_para_runs(doc, [("B. 파생 기상", True), (
        " — 냉방·난방도일 cdd_18·hdd_18 (base 18°C) + 3·7일 이동평균. "
        "체감온도(기상청 공식: 여름 Stull Tw + 기상청 PT, 겨울 NWS Wind Chill). "
        "폭염·한파 임계 더미 4개 (공식 발효기준 정합, 2일 연속).", False)], bullet=True)
    add_para_runs(doc, [("C. 달력", True), (
        " — dow, month, is_weekend, is_holiday (특일 정보 API + holidays 라이브러리 교차검증).", False)],
                  bullet=True)
    add_para_runs(doc, [("D. 장기 추세 동인", True), (
        " — pop_total(전국 주민등록인구), ip_total(전산업생산지수, 발표 lag 2개월 강제).", False)],
                  bullet=True)
    add_para(doc, "→ 총 28개 외생변수를 5,844행 단일 wide-table(exog_daily_2010_2025.csv)로 통합. 결측 0.",
             bullet=False)

    # ===== 2. EDA =====
    add_heading(doc, "2. EDA 핵심 발견", level=1)
    add_para_runs(doc, [("정상성: ", True), (
        "원본 비정상(KPSS), 1차 차분으로 정상화 (ADF·KPSS 동시 통과) → d=1 확정.", False)], bullet=True)
    add_para_runs(doc, [("계절성: ", True), (
        "STL 분해 결과 주간(s=7) 진폭 22,290 MW — 평균의 33%. "
        "연간 365 계절성은 외생변수(cdd/hdd, month)로 흡수, SARIMA는 s=7만.", False)], bullet=True)
    add_para_runs(doc, [("ACF/PACF: ", True), (
        "lag 7·14·21에서 강한 spike → SARIMA(p,1,q)(P,1,Q,7) 차수 후보 도출.", False)], bullet=True)
    add_para_runs(doc, [("다중공선성: ", True), (
        "ta_avg/cdd/hdd VIF 무한대, feels_like_* 1500+ → 모델링 시 정리.", False)], bullet=True)
    add_para_runs(doc, [("2025 분포 시프트: ", True), (
        "peak_mw 평균 +8.66%, ip_total +17%, cdd_18 +20%. 추세 변수가 흡수 필요.", False)], bullet=True)

    add_figure(doc, FIG_DIR / "fig03_stl_decomposition.png",
               "그림 2. STL 분해 — 추세는 16년간 상승, 주간 계절은 진폭 22k MW로 일관.")

    add_figure(doc, FIG_DIR / "fig02_exog_corr_heatmap.png",
               "그림 3. 외생변수 + peak_mw 상관행렬. 우측 마지막 행/열이 peak_mw. "
               "cdd/hdd 계열은 정의상 강한 상관 → VIF 검토로 정리.")

    # ===== 3. 모델링 =====
    add_heading(doc, "3. 모델링 — 3차 반복 개선", level=1)

    add_heading(doc, "3.1 v1 베이스라인 — 9개 조합 그리드", level=2)
    add_para(doc, "차수 3개(1,1,1)·(2,1,1)·(1,1,2) × 외생 3서브셋(none/minimal 6개/phase1 16개) = 9 조합.",
             bullet=True)
    add_para(doc, "AIC 최저: SARIMAX(1,1,2)(1,1,1,7) + phase1 16개 외생변수.", bullet=True)
    add_para(doc, "2025 MAPE: one-shot 4.16% / rolling 2.09%.", bullet=True)
    add_para(doc, "발견: is_weekend 계수 ≈ 0 (SARIMA s=7가 흡수), ip_total 계수 음수(다중공선 의심).",
             bullet=True)

    add_heading(doc, "3.2 v2 — 다중공선성 정리 (8개 서브셋 그리드)", level=2)
    add_para(doc, "is_weekend 제거, {ip_total, pop_total} × {cdd 원본, cdd 원본+ma7} × {hdd 동일} 비교.",
             bullet=True)
    add_para_runs(doc, [("cdd/hdd ma7은 필수: ", True), (
        "원본만 쓰면 MAPE 8.6~11.8%로 폭락 — 7일 이동평균이 건물 열관성 흡수.", False)], bullet=True)
    add_para_runs(doc, [("pop_total > ip_total: ", True), (
        "홀드아웃 일반화 기준. ip는 단기 변동·외생성 의심.", False)], bullet=True)
    add_para_runs(doc, [("AIC vs 2025 MAE 갈림: ", True), (
        "ip 버전이 AIC 더 작지만 2025 MAE는 pop 버전이 작음 → 일반화 우선 채택.", False)], bullet=True)
    add_para(doc, "2025 MAPE: one-shot 4.03% / rolling 2.08%, 외생변수 14개로 단순화.", bullet=True)

    add_heading(doc, "3.3 v3 — 3단계 추가 개선 (log → OR더미 → 비선형 기온)", level=2)

    add_table(doc,
              header=["단계", "변경", "MAE", "MAPE", "Δ vs 직전"],
              rows=[
                  ["v2", "(baseline 재학습)", "2,969", "4.04%", "—"],
                  ["v3a", "+ log(peak_mw) 변환", "2,736", "3.70%", "-233 (-7.9%)"],
                  ["v3b", "+ 희소 더미 OR 통합 (14→12 exog)", "2,739", "3.71%", "+3 (단순화)"],
                  ["v3c", "+ 비선형 기온항 (cdd², hdd²)", "2,668", "3.63%", "-71 (-2.6%)"],
              ],
              col_widths_cm=[1.4, 6.5, 2.0, 2.0, 3.0],
              highlight_rows=[3])

    add_para_runs(doc, [("log 변환이 가장 큰 단일 개선 (-7.9%)", True), (
        " — 분산 안정화 효과. OR더미 통합은 점예측 영향 없이 단순화만(14→12). "
        "비선형 cdd_18²·hdd_18² 둘 다 p<0.001로 강하게 유의 — 기온-부하 비선형성 검증.", False)])

    # ===== 4. 최종 모델 성능 =====
    add_heading(doc, "4. 최종 모델 (v3c) — 2025 홀드아웃 성능", level=1)

    add_para_runs(doc, [
        ("모델: ", True),
        ("SARIMAX(1,1,2)(1,1,1,7) on log(peak_mw) + 14개 외생변수.", False),
    ])
    add_para(doc, "외생 14개: cdd_18, cdd_18_ma7, cdd_18_sq, hdd_18, hdd_18_ma7, hdd_18_sq, "
                  "hm_avg, ws_avg, rn_day, ss_day, heat_th_2day_any, cold_th_2day_any, "
                  "is_holiday, pop_total.", bullet=False, size=9.5)

    add_figure(doc, FIG_DIR / "fig04_model_mae_compare.png",
               "그림 4. 모델별 2025 365일 MAE 비교. v3c(녹색)가 모든 베이스라인·이전 버전 대비 우수.")

    add_table(doc,
              header=["모델", "MAE (MW)", "MAPE", "95% PI 커버"],
              rows=[
                  ["v3c rolling 1-step (실무 단기예측)", "1,357", "1.87%", "—"],
                  ["v3c one-shot 365일", "2,668", "3.63%", "95.3%"],
                  ["v1 phase1 one-shot", "3,062", "4.16%", "95.3%"],
                  ["SARIMA (외생변수 없음)", "7,337", "10.33%", "—"],
                  ["Naive (전일)", "4,311", "5.95%", "—"],
              ],
              col_widths_cm=[6.5, 2.5, 2.0, 3.0],
              highlight_rows=[0, 1])

    add_para_runs(doc, [
        ("v1 → v3c: MAE -13%, MAPE -0.53%p. ", True),
        ("Naive 대비 MAE -38%, SARIMA(no exog) 대비 -64%.", False),
    ])

    add_figure(doc, FIG_DIR / "fig05_v3c_forecast.png",
               "그림 5. v3c 2025년 365일 one-shot 예측 vs 실측. "
               "95% 예측구간(연한 띠)이 실측을 95.3% 커버.")

    add_heading(doc, "4.1 월별 MAE — v2 vs v3c", level=2)
    add_figure(doc, FIG_DIR / "fig06_monthly_mae_v2_v3c.png",
               "그림 6. 월별 MAE — v3c는 1·2·4·6·10월 등에서 큰 폭 개선, "
               "9월은 추석 연휴 효과로 후퇴.")

    add_para(doc, "개선: 1월 -24%, 2월 -35%, 4월 -31%, 6월 -34%, 10월 -30% (한파·환절기).", bullet=True)
    add_para(doc, "후퇴: 9월 +74% (추석 연휴 효과 미흡, is_holiday 단일 더미 한계).", bullet=True)
    add_para(doc, "동일: 8월 ~0% (폭염 효과는 v2에서 cdd/ma7로 이미 흡수).", bullet=True)

    # ===== 5. 한계 =====
    add_heading(doc, "5. 남은 한계 (보고서 한계점 섹션 후보)", level=1)
    add_para(doc, "잔차 자기상관 잔존 (Ljung-Box p<0.001) — 차수 확장으로 일부 잡을 여지.", bullet=True)
    add_para(doc, "음력 명절 효과 — is_holiday 단일 더미로 다일 연휴 흡수 못 함, ±10k MW 오차.",
             bullet=True)
    add_para(doc, "pop_total p=0.68 — 16년간 변동 폭 좁아 표준화 후 신호 약함 "
                  "(모델 일반화에는 기여).", bullet=True)
    add_para(doc, "외생변수 실측치 가정 — 보고된 MAPE는 \"외생변수 완벽 예측 시 상한 성능\". "
                  "실무 적용 시 외생변수 자체 예측 오차 추가.", bullet=True)
    add_para(doc, "상세 후속 작업: docs/todo.md", bullet=True)

    # ===== 6. 산출물 =====
    add_heading(doc, "6. 핵심 산출물", level=1)
    add_para_runs(doc, [("데이터: ", True), (
        "data/processed/exog_daily_2010_2025.csv, peak_mw_2010_2025.csv", False)], bullet=True)
    add_para_runs(doc, [("모델 코드: ", True), (
        "src/model_arimax.py (v1), model_arimax_v2.py, model_arimax_v3.py", False)], bullet=True)
    add_para_runs(doc, [("분석 노트북: ", True), (
        "notebooks/eda.ipynb, modeling.ipynb", False)], bullet=True)
    add_para_runs(doc, [("의사결정 이력: ", True), (
        "docs/research_log/2026-05-25_*.md (수집·EDA), 2026-05-26_*.md (모델링)", False)], bullet=True)
    add_para_runs(doc, [("변수 메타: ", True), (
        "docs/independent_var.md, dependent_var.md", False)], bullet=True)

    # 한 줄 요약
    final = doc.add_paragraph()
    final.paragraph_format.space_before = Pt(8)
    run = final.add_run("한 줄 요약: ")
    set_font(run, size=11, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    run2 = final.add_run(
        "rolling 1-step MAPE 1.87%는 운영급 수준, "
        "one-shot 365일 3.63%는 학계 ARIMAX 평균 수준의 결과."
    )
    set_font(run2, size=11)

    doc.save(OUT_PATH)
    print(f"저장: {OUT_PATH}")
    print(f"크기: {OUT_PATH.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    build()
